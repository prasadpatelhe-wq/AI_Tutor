import os
import re
import tempfile
from fastapi import APIRouter, UploadFile, File, HTTPException
from docx import Document
import fitz  # PyMuPDF

router = APIRouter(prefix="/syllabus", tags=["Syllabus"])


def extract_text_from_docx(file_path):
    """Extract text from DOCX file."""
    doc = Document(file_path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


def extract_text_from_pdf(file_path):
    """Extract clean text from PDF, handling layout issues in CBSE/NCERT files."""
    import fitz  # PyMuPDF

    text_blocks = []
    with fitz.open(file_path) as pdf:
        for page in pdf:
            # Use layout=True for better block separation
            blocks = page.get_text("blocks")
            blocks = sorted(blocks, key=lambda b: (b[1], b[0]))  # sort top-to-bottom, left-to-right
            for b in blocks:
                line = b[4].strip()
                if line:
                    text_blocks.append(line)

    # Join blocks as lines
    text = "\n".join(text_blocks)

    # Clean up spacing
    text = text.replace("  ", " ")
    text = text.replace("\t", " ").strip()

    return text



def extract_chapters_from_text(text):
    """
    Extracts top-level chapters like '1. Real Numbers' or '1 Real Numbers',
    ignoring sub-sections like '3.1 Substitution Method'.
    Works for CBSE, NCERT, and similar formats.
    """
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    chapters = []

    for line in lines:
        line_clean = line.replace("–", "-").replace("—", "-")

        # ✅ Matches:
        # 1 Real Numbers
        # 1. Real Numbers
        # Chapter 1: Real Numbers
        # Unit 2 - Polynomials
        match = re.match(r'^(?:Chapter|Unit)?\s*\d+\.?\s*[-:]?\s*([A-Za-z][A-Za-z0-9\s\-&():]+)$', line_clean, re.IGNORECASE)

        if match:
            title = match.group(1).strip()

            # Skip subsection lines like 3.1, 4.2, etc.
            if not re.match(r'^\d+\.\d+', line_clean):
                chapters.append({
                    "id": str(len(chapters) + 1),
                    "title": title,
                    "summary": f"This chapter covers {title} in detail."
                })

    if not chapters:
        raise HTTPException(status_code=400, detail="No chapters detected in syllabus file. Try uploading a clean CBSE syllabus PDF.")

    # Deduplicate
    seen = set()
    final = []
    for ch in chapters:
        if ch["title"] not in seen:
            final.append(ch)
            seen.add(ch["title"])

    print(f"✅ Extracted {len(final)} chapters.")
    return final

@router.post("/upload")
async def upload_syllabus(file: UploadFile = File(...)):
    """
    Upload syllabus file (PDF/DOCX) and extract CBSE-style numbered chapters.
    Example: '1. Real Numbers', '2. Polynomials', etc.
    """
    import re
    import fitz
    from docx import Document
    import tempfile, os

    def extract_text_from_pdf(file_path):
        """Extract clean text from PDF."""
        text = ""
        with fitz.open(file_path) as pdf:
            for page in pdf:
                text += page.get_text("text")
        return text

    try:
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as tmp:
            tmp.write(await file.read())
            tmp_path = tmp.name

        ext = os.path.splitext(file.filename)[1].lower()
        if ext == ".docx":
            text = "\n".join([p.text for p in Document(tmp_path).paragraphs if p.text.strip()])
        elif ext == ".pdf":
            text = extract_text_from_pdf(tmp_path)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")

        os.remove(tmp_path)

        # --- Extract only "1. Chapter Name" style entries ---
        pattern = r'(?m)^\s*(\d+)\.\s*([A-Za-z].+?)\s+\d+$'
        matches = re.findall(pattern, text)

        chapters = []
        for num, title in matches:
            title = title.strip()
            # Stop if we reach appendix or answers
            if re.search(r'Appendix|Answers|Hints', title, re.IGNORECASE):
                break
            chapters.append({
                "id": num,
                "title": title,
                "summary": f"This chapter covers {title} in detail."
            })

        if not chapters:
            raise HTTPException(status_code=400, detail="No chapters detected in syllabus file.")

        print("\n✅ Extracted Chapters:")
        for ch in chapters:
            print(f"- {ch['title']}")

        return {"message": "✅ Syllabus processed successfully", "chapters": chapters}

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing syllabus: {str(e)}")
